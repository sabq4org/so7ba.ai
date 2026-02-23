#!/usr/bin/env python3
"""Build Word document for 1st Intermediate English Exam (Super Goal 1) - Second Term 1447"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = 'Times New Roman'
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

def add_heading_centered(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def add_bold_para(text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p

def add_para(text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def add_line(text):
    return doc.add_paragraph(text)

def add_exam_header(school_blank, exam_title, grade, term_info, score_text):
    t = doc.add_table(rows=3, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.cell(0,0).text = "Kingdom of Saudi Arabia"
    t.cell(1,0).text = "Ministry of Education"
    t.cell(2,0).text = school_blank
    t.cell(0,1).text = "وزارة التعليم\nMinistry of Education"
    t.cell(1,1).text = exam_title
    t.cell(0,2).text = grade
    t.cell(1,2).text = term_info
    t.cell(2,2).text = score_text
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def section_break():
    doc.add_page_break()

def add_separator():
    doc.add_paragraph("─" * 60)

# ============================================================
# EXAM 1: Quiz (1) - Ms. Tahani Al-Ghamdi (Pages 2-5)
# ============================================================
add_heading_centered("Quiz (1) – 1447 – Second Term", level=1)
add_heading_centered("1st Intermediate Grade", level=2)

add_exam_header(
    "…… Intermediate School",
    "Quiz (1) - 1447",
    "Name: ……………………………",
    "Class: 1st Intermediate Grade",
    "Second Term    10 /"
)

# I. Comprehension
add_bold_para("I. Comprehension:")
add_bold_para("Read the passage to answer the questions below:")
add_para("Max is my little cat. He is one year old. He is brown. He is very cute. He likes to play around the house. He loves to play with me. He is very smart. His favourite place to sleep is under my desk, but sometimes he sleeps in the kitchen. He loves to eat hot dogs a lot. He is a very good and funny cat. I love him.")
add_para("")
add_bold_para("A- Choose the correct answer:")
add_line("1- This text is about ………………")
add_line("    a- The writer's family       b- The writer's cat       c- The writer's dog")
add_line("2- Max's color is ………..")
add_line("    a- black       b- grey       c- brown")
add_para("")
add_bold_para("B- Check ( T ) or ( F ):")
add_line("1- Max's favourite food is burger. (     )")
add_line("2- Max likes to sleep under the desk. (     )")

# II. Grammar
add_separator()
add_bold_para("II. Grammar:")
add_bold_para("A- Choose:")
add_line("1- I like coffee (and – but – or) I don't like tea.")
add_line("2- Ali and Omar (play – plays – playing) tennis.")
add_line("3- Leen (doesn't watches – doesn't watch – don't watch) TV every night.")
add_line("4- (Do – Does – Are) they teach math?")
add_para("")
add_bold_para("B- Do as shown between brackets:")
add_line("1- Yara usually (write) a letter to her friend.       (Correct the mistake)")
add_line("……………………………………………………………")
add_line("2- has / black / Sara / hair       (Rearrange)")
add_line("…………………………………………………….")

# III. Vocabulary
add_separator()
add_bold_para("III. Vocabulary:")
add_bold_para("A- Classify the words:")
add_line("thin – history – clinic")
t = doc.add_table(rows=2, cols=3)
t.style = 'Table Grid'
t.cell(0,0).text = "School subjects"
t.cell(0,1).text = "Adjectives to describe people's look"
t.cell(0,2).text = "Place of work"
for row in t.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

section_break()
add_bold_para("B- Write the words under the pictures:")
add_line("carpenter – lawyer – confused – experiment")
add_line("[Picture 1: lawyer]    [Picture 2: experiment]    [Picture 3: carpenter]")
add_line("______________       ______________       ______________")

# IV. Orthography
add_separator()
add_bold_para("IV. Orthography:")
add_bold_para("A- Choose the correct missing letters:")

t = doc.add_table(rows=4, cols=4)
t.style = 'Table Grid'
data = [
    ["1- He is an a__tive student.", "a- c", "b- k", "c- x"],
    ["2- She has b__ond hair and blue eyes.", "a- l", "b- r", "c- c"],
    ["3- These stores s__ll clothes and shoes.", "a- a", "b- e", "c- i"],
    ["4- Ali is not old enough to dri__e a car.", "a- f", "b- v", "c- p"],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        t.cell(i,j).text = val

add_para("")
add_bold_para("Good Luck ☺", WD_ALIGN_PARAGRAPH.CENTER)
add_bold_para("Ms. Tahani Al-Ghamdi", WD_ALIGN_PARAGRAPH.CENTER)

# ---- ANSWER KEY ----
section_break()
add_heading_centered("نموذج الإجابة – Quiz (1)", level=1)
add_heading_centered("Answer Key", level=2)

add_bold_para("I. Comprehension:")
add_line("A- 1: b- The writer's cat")
add_line("A- 2: c- brown")
add_line("B- 1: F")
add_line("B- 2: T")
add_para("")
add_bold_para("II. Grammar:")
add_line("A- 1: but")
add_line("A- 2: play")
add_line("A- 3: doesn't watch")
add_line("A- 4: Do")
add_line("B- 1: Yara usually writes a letter to her friend.")
add_line("B- 2: Sara has black hair.")
add_para("")
add_bold_para("III. Vocabulary:")
add_line("A- School subjects: history | Adjectives: thin | Place of work: clinic")
add_line("B- lawyer / experiment / carpenter")
add_para("")
add_bold_para("IV. Orthography:")
add_line("1- c (active)")
add_line("2- l (blond)")
add_line("3- e (sell)")
add_line("4- v (drive)")

# ============================================================
# EXAM 2: 1st Period Exam (Pages 6-9)
# ============================================================
section_break()
add_heading_centered("1st Period Exam – 1447 – Second Term", level=1)
add_heading_centered("1st Intermediate Grade – English Language     /20", level=2)

add_line("Student's Name: ……………………………………    Class: ……………")

# 1- Reading
add_bold_para("1- Read the following passage and answer the questions below:    /5 marks")
add_para("Omar is a doctor who works at a hospital in Saudi Arabia. He helps sick people every day and always listens carefully to his patients. Omar wears a white coat and works long hours, but he loves his job. He feels happy when his patients feel better and go home smiling.")
add_para("")
add_bold_para("A. Choose the correct answer:")

t = doc.add_table(rows=3, cols=4)
t.style = 'Table Grid'
t.cell(0,0).text = "1- What is Omar's job?"
t.cell(0,1).text = "a- teacher"
t.cell(0,2).text = "b- engineer"
t.cell(0,3).text = "c- doctor"
t.cell(1,0).text = "2- Where does Omar work?"
t.cell(1,1).text = "a- museum"
t.cell(1,2).text = "b- hospital"
t.cell(1,3).text = "c- school"
t.cell(2,0).text = "3- What does Omar wear at work?"
t.cell(2,1).text = "a- T-shirt"
t.cell(2,2).text = "b- a white coat"
t.cell(2,3).text = "c- pants"

add_para("")
add_bold_para("B. Put ( ✓ ) or ( ✗ ):")
add_line("1- Omar helps sick people every day. (        )")
add_line("2- He feels sad when his patients feel better. (        )")

# 2- Grammar
add_separator()
add_bold_para("2- Choose the correct answer:    /5 marks")

t = doc.add_table(rows=5, cols=4)
t.style = 'Table Grid'
t.cell(0,0).text = "1- I ………work in school."
t.cell(0,1).text = "a- am not"
t.cell(0,2).text = "b- doesn't"
t.cell(0,3).text = "c- don't"
t.cell(1,0).text = "2- Faisal ………….a red backpack."
t.cell(1,1).text = "a- have"
t.cell(1,2).text = "b- has"
t.cell(1,3).text = "c- is"
t.cell(2,0).text = "3- He's a bus driver. He………. a bus."
t.cell(2,1).text = "a- drive"
t.cell(2,2).text = "b- drives"
t.cell(2,3).text = "c- driving"
t.cell(3,0).text = "4- She is really ……. about the news."
t.cell(3,1).text = "a- depress"
t.cell(3,2).text = "b- depressing"
t.cell(3,3).text = "c- depressed"
t.cell(4,0).text = "5- Nawal is ……….her homework at the moment."
t.cell(4,1).text = "a- do"
t.cell(4,2).text = "b- does"
t.cell(4,3).text = "c- doing"

# 3- Do as shown
add_separator()
add_bold_para("3- Do as shown between brackets:    /2 marks")
add_bold_para("A. Correct:")
add_line("1- He wants to be a carpenter so he likes to build houses……………")
add_bold_para("B. Reorder:")
add_line("1- Math / quite / a difficult / is / subject")
add_line("……………………………………………………………………………")

section_break()

# 4- Pictures
add_bold_para("4- Write the correct word under each picture:    /4 marks")
add_line("excited – carpenter – night – driver")
add_line("[Picture 1: night scene]    [Picture 2: driver]    [Picture 3: carpenter]    [Picture 4: excited person]")
add_line("…………………    …………………    …………………    …………………")

# 5- Choose
add_separator()
add_bold_para("5- Choose the correct answer:    /2 marks")
add_line("1- Members learn lots of computer programs in (computer – science – archaeology) Club.")
add_line("2- Ahmad's brother works in a hospital. He's a (reporter – flight attendant – doctor)")

# 6- Missing letter choose
add_separator()
add_bold_para("6- Choose the missing letter:    /2 marks")
add_line("1- He has brown eyes and …urly black hair. ( u – c – w )")
add_line("2- I want to be a pilo… because I'm interested in planes. ( f – t – r )")

# 7- Write missing letter
add_separator()
add_bold_para("7- Write the missing letter:")
add_line("[Picture: Bored person]    Bo……ed")
add_line("[Picture: Photographer]    ……hotographer")

add_bold_para("GOOD LUCK", WD_ALIGN_PARAGRAPH.CENTER)

# ---- ANSWER KEY ----
section_break()
add_heading_centered("نموذج الإجابة – 1st Period Exam", level=1)
add_line("1- A: 1-c (doctor), 2-b (hospital), 3-b (a white coat)")
add_line("   B: 1- ✓, 2- ✗")
add_line("2- 1-c (don't), 2-b (has), 3-b (drives), 4-c (depressed), 5-c (doing)")
add_line("3- A: because (instead of so)")
add_line("   B: Math is quite a difficult subject")
add_line("4- night / driver / carpenter / excited")
add_line("5- 1: computer, 2: doctor")
add_line("6- 1: c (curly), 2: t (pilot)")
add_line("7- Bored (r), Photographer (p)")

# ============================================================
# EXAM 3: First Period Exam - With Oral/Written marks (Pages 10-13)
# ============================================================
section_break()
add_heading_centered("First Period Exam – Second Term", level=1)
add_heading_centered("First Intermediate Grade", level=2)

t = doc.add_table(rows=2, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.cell(0,0).text = "المجموع\nTotal"
t.cell(0,1).text = "شفهي"
t.cell(0,2).text = "تحريري"
t.cell(1,0).text = "20"
t.cell(1,1).text = "10"
t.cell(1,2).text = "10"
for row in t.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_line("Name …………………………………          Class ………………………………")

# Comprehension
add_separator()
add_bold_para("Comprehension    /2")
add_bold_para("Read the following then answer the questions:")
add_para("Omar Hamdan lives in Tabuk. He is sixteen years old, and he's on the school football team. Omar is a very good player, and he's the team's top striker this season. Omar wants to be a professional football player. His parents support him, but they want him to go to a university.")

add_bold_para("A- Put (✓) for true or (X) for false:")
add_line("1. Omar Hamdan is on the school football team. (          )")
add_line("2. He wants to be a teacher. (          )")

add_bold_para("B- Choose the right answer:")
t = doc.add_table(rows=2, cols=5)
t.style = 'Table Grid'
t.cell(0,0).text = "1- Omar lives in ______."
t.cell(0,1).text = "Riyadh"
t.cell(0,2).text = "Jeddah"
t.cell(0,3).text = "Abha"
t.cell(0,4).text = "Tabuk"
t.cell(1,0).text = "2- His parents want him to go to a _____."
t.cell(1,1).text = "school"
t.cell(1,2).text = "university"
t.cell(1,3).text = "broad"
t.cell(1,4).text = "market"

# Grammar
add_separator()
add_bold_para("Grammar    /2.5")
add_bold_para("Choose the right answer:")

t = doc.add_table(rows=5, cols=5)
t.style = 'Table Grid'
t.cell(0,0).text = "1- He ______ for an airline."
t.cell(0,1).text = "A. work"
t.cell(0,2).text = "B. works"
t.cell(0,3).text = "C. working"
t.cell(0,4).text = "D. do"
t.cell(1,0).text = "2- _____ does he do? He is a bus driver."
t.cell(1,1).text = "A. What"
t.cell(1,2).text = "B. when"
t.cell(1,3).text = "C. Where"
t.cell(1,4).text = "D. Which"
t.cell(2,0).text = "3- They ____ English."
t.cell(2,1).text = "A. speakes"
t.cell(2,2).text = "B. speaks"
t.cell(2,3).text = "C. speak"
t.cell(2,4).text = "D. speaking"
t.cell(3,0).text = "4- Do you watch the film? Yes, I ________."
t.cell(3,1).text = "A. doesn't"
t.cell(3,2).text = "B. does"
t.cell(3,3).text = "C. don't"
t.cell(3,4).text = "D. do"
t.cell(4,0).text = "5- ______ is he taking a nap? Because he's tired."
t.cell(4,1).text = "A. Why"
t.cell(4,2).text = "B. What"
t.cell(4,3).text = "C. Who"
t.cell(4,4).text = "D. How"

# Vocabulary
section_break()
add_bold_para("Vocabulary    /2")
add_bold_para("A- Choose the best answer to fill in the blank:")

t = doc.add_table(rows=4, cols=5)
t.style = 'Table Grid'
t.cell(0,0).text = "1. She wants to be a _______. She wants to help sick people."
t.cell(0,1).text = "nurse"
t.cell(0,2).text = "teacher"
t.cell(0,3).text = "pilot"
t.cell(0,4).text = "designer"
t.cell(1,0).text = "2. I want to be a flight attendant. I like to _____ and meet lots of different people."
t.cell(1,1).text = "sleep"
t.cell(1,2).text = "travel"
t.cell(1,3).text = "play"
t.cell(1,4).text = "study"
t.cell(2,0).text = "3. What's your favorite ______, Carl? I prefer history."
t.cell(2,1).text = "month"
t.cell(2,2).text = "week"
t.cell(2,3).text = "subject"
t.cell(2,4).text = "school"
t.cell(3,0).text = "4- He _________ things clearly and also makes math fun."
t.cell(3,1).text = "helps"
t.cell(3,2).text = "sleeps"
t.cell(3,3).text = "opens"
t.cell(3,4).text = "explains"

add_para("")
add_bold_para("B- Put the right word under its picture:    /2.5")
add_line("smartphone – tired – science – waiter – cameraman")
add_line("[Picture 1: tired]  [Picture 2: science]  [Picture 3: smartphone]  [Picture 4: cameraman]  [Picture 5: waiter]")
add_line("______________  ______________  ______________  ______________  ______________")

# Orthography
add_separator()
add_bold_para("Orthography    /1")
add_bold_para("Choose the correct letter:")
t = doc.add_table(rows=2, cols=5)
t.style = 'Table Grid'
t.cell(0,0).text = "1."
t.cell(0,1).text = "My mother teaches at a pri_ary school."
t.cell(0,2).text = "l"
t.cell(0,3).text = "m"
t.cell(0,4).text = "s"
t.cell(1,0).text = "2."
t.cell(1,1).text = "Matt has blon__ hair."
t.cell(1,2).text = "t"
t.cell(1,3).text = "f"
t.cell(1,4).text = "d"

add_bold_para("Good luck          T.", WD_ALIGN_PARAGRAPH.CENTER)

# ---- ANSWER KEY ----
section_break()
add_heading_centered("نموذج الإجابة – First Period Exam (Oral/Written)", level=1)
add_line("Comprehension: A- 1: ✓, 2: X | B- 1: Tabuk, 2: university")
add_line("Grammar: 1-B (works), 2-A (What), 3-C (speak), 4-D (do), 5-A (Why)")
add_line("Vocabulary: A- 1: nurse, 2: travel, 3: subject, 4: explains")
add_line("           B- tired / science / smartphone / cameraman / waiter")
add_line("Orthography: 1: m (primary), 2: d (blond)")

# ============================================================
# EXAM 4: Fahad Basketball Exam (Pages 14-17)
# ============================================================
section_break()
add_heading_centered("1st Period Exam – English Language", level=1)
add_heading_centered("1st Intermediate Grade", level=2)

add_line("Name: ……………………………………    Class: ……………    Mark: ………………")

add_bold_para("❊ Reading Comprehension ❊    /5 marks")
add_para("Fahad lives in Jeddah. He is fifteen years old. He plays basketball at school. Fahad is a very good player. He is the best defender on his team.\n\nFahad wants to join a big basketball club one day. His family supports him, but they also want him to study well. Fahad is happy because he will go to a Basketball Camp this winter to learn new skills.")

add_bold_para("1) Fill in the blank with the correct words:")
t = doc.add_table(rows=3, cols=4)
t.style = 'Table Grid'
t.cell(0,0).text = "1. Where does Fahad live?"
t.cell(0,1).text = "a. Riyadh"
t.cell(0,2).text = "b. Jeddah"
t.cell(0,3).text = "c. Tabuk"
t.cell(1,0).text = "2. How old is Fahad?"
t.cell(1,1).text = "a. 14 years"
t.cell(1,2).text = "b. 15 years"
t.cell(1,3).text = "c. 16 years"
t.cell(2,0).text = "3. What sport does Fahad play?"
t.cell(2,1).text = "a. basketball"
t.cell(2,2).text = "b. football"
t.cell(2,3).text = "c. tennis"

add_bold_para("2) Choose (✓) or (✗):")
add_line("1. Fahad plays basketball at school. (  )")
add_line("2. His family does not support him. (  )")

add_separator()
add_bold_para("❊ Grammar ❊    /7 marks")
add_bold_para("Choose the correct:")
t = doc.add_table(rows=5, cols=4)
t.style = 'Table Grid'
t.cell(0,0).text = "1. My father __________ in a bank."
t.cell(0,1).text = "worked"
t.cell(0,2).text = "work"
t.cell(0,3).text = "works"
t.cell(1,0).text = "2. We always __________ lunch in the living room."
t.cell(1,1).text = "had"
t.cell(1,2).text = "have"
t.cell(1,3).text = "has"
t.cell(2,0).text = "3. Why are you late? ________ I missed the bus."
t.cell(2,1).text = "because"
t.cell(2,2).text = "but"
t.cell(2,3).text = "so"
t.cell(3,0).text = "4. I get up... ___six o'clock."
t.cell(3,1).text = "on"
t.cell(3,2).text = "at"
t.cell(3,3).text = "in"
t.cell(4,0).text = "5. History is a very __________ subject."
t.cell(4,1).text = "bore"
t.cell(4,2).text = "bored"
t.cell(4,3).text = "boring"

add_para("")
add_bold_para("Write the correct answer:")
add_line("1. Does you like math? ___________ (correct)")
add_line("2. long / hair / has / blond / she _________________________ (order)")

section_break()
add_bold_para("❊ Vocabulary ❊    /6 marks")
add_bold_para("1) Match the following:")
t = doc.add_table(rows=4, cols=2)
t.style = 'Table Grid'
t.cell(0,0).text = "1. awesome"
t.cell(0,1).text = "difficult"
t.cell(1,0).text = "2. hard"
t.cell(1,1).text = "really great"
t.cell(2,0).text = "3. designs new gadgets"
t.cell(2,1).text = "tennis player"
t.cell(3,0).text = "4. plays tennis"
t.cell(3,1).text = "hi-tech designer"

add_para("")
add_bold_para("2) Choose the correct answer:")
add_line("1. Mustafa is not short. He's ______________. (tall – athletic – smart)")
add_line("2. Ahmed drives a taxi. He's a ______________. (taxi driver – player – lawyer)")
add_line("3. Nura likes to help others. She's ______________. (thin – friendly – short)")
add_line("4. Mustafa sells cars. He's a ______________. (doctor – teacher – salesperson)")

add_para("")
add_bold_para("3) Write the correct words above each picture:")
add_line("bored / smartphone / clock / mechanic")
add_line("[Picture 1: clock]  [Picture 2: mechanic]  [Picture 3: smartphone]  [Picture 4: bored person]")
add_line("______________  ______________  ______________  ______________")

add_separator()
add_bold_para("❊ Orthography ❊    /2 marks")
add_bold_para("Choose the correct letter:")
add_line("1. A doc__or works in a hospital. (d – t – k)")
add_line("2. My favorite subject is hist_ry. (n – d – o)")
add_line("3. I use a comp_ter to play games. (u – o – q)")
add_line("4. He joined a dra__a club. (n – m – u)")

add_bold_para("☺ With my best wishes ☺", WD_ALIGN_PARAGRAPH.CENTER)
add_bold_para("Focus with Manal          Teacher/Manal Al-Shareef", WD_ALIGN_PARAGRAPH.CENTER)

# ---- ANSWER KEY ----
section_break()
add_heading_centered("نموذج الإجابة – Fahad Basketball Exam", level=1)
add_line("Reading: 1-b (Jeddah), 2-b (15 years), 3-a (basketball) | ✓, ✗")
add_line("Grammar: 1-works, 2-have, 3-because, 4-at, 5-boring")
add_line("         1: Do you like math? | 2: She has long blond hair.")
add_line("Vocabulary: Match: 1→2, 2→1, 3→4, 4→3")
add_line("           Choose: 1-tall, 2-taxi driver, 3-friendly, 4-salesperson")
add_line("           Pictures: clock / mechanic / smartphone / bored")
add_line("Orthography: 1-t (doctor), 2-o (history), 3-u (computer), 4-m (drama)")

# ============================================================
# EXAM 5: Quiz 1 - Rania (Pages 18-19)
# ============================================================
section_break()
add_heading_centered("Quiz 1 – Units 9+10 – Super Goal 1", level=1)
add_heading_centered("First Intermediate", level=2)

add_line("Name: ……………………………    /10    Class: …………")

add_bold_para("I – Comprehension    /2")
add_bold_para("Read then answer the questions:")
add_para("My younger sister Rania has a weight problem because she eats lots of sweets. My mother took her to the doctor. He asked her, \"How much sugar do you put in a glass of milk?\" She said, \"Two spoons.\" He told her, \"Rania, You shouldn't take so much sugar. Sugar is an 'empty' food. It gives you energy quickly, but it doesn't help you stay healthy. I know it's not easy to stop taking sugar, but you can try it. Eat fruit instead of biscuits, cake and chocolates. You'll soon become healthy.\"")

add_bold_para("Choose the correct answer:")
add_line("1- Rania has a weight problem because she eats less food. A- Yes  B- No")
add_line("2- Rania put ………………………… spoons of sugar in her milk.")
add_line("   A- two    B- three    C- four")

add_separator()
add_bold_para("II. Grammar    /3")
add_bold_para("A – Choose the correct answers:")
add_line("3- Ahmed is a doctor. He ………… in a hospital.")
add_line("   A- work    B- works    C- working")
add_line("4- What ……… you want to be?")
add_line("   A- do    B- does    C- are")
add_line("5- Hind wants to be a doctor ……………… she wants to help sick people.")
add_line("   A- so    B- because    C- and")
add_line("6- Ali doesn't ………….. French.")
add_line("   A- study    B- studies    C- studying")
add_line("7- A: Does Tom teach geography? B: Yes, ……….")
add_line("   A- they do    B- he do    C- he does")
add_line("8- Leen ……………………… wear glasses.")
add_line("   A- doesn't    B- do    C- don't")

section_break()
add_bold_para("III. Vocabulary    /3")
add_bold_para("A- Write the words under the pictures:")
add_line("A- teacher    B- nurse    C- carpenter    D- cyclist")
add_line("[Picture 9]  ______________    [Picture 10]  ______________")

add_bold_para("B- Write the correct words under each picture:")
add_line("A- Science    B- Computer science    C- Art    D- History")
add_line("[Picture 11]  ______________    [Picture 12]  ______________")

add_bold_para("C- Circle the odd word:")
add_line("13- A- geography    B- PE    C- team    D- science")
add_line("14- A- easy    B- fun    C- challenging    D- Art")

add_separator()
add_bold_para("IV. Spelling    /2")
add_bold_para("Choose the words with the correct spelling:")
t = doc.add_table(rows=2, cols=4)
t.style = 'Table Grid'
t.cell(0,0).text = "15-\nA- mat  B- Math\nC- math  D- meth"
t.cell(0,1).text = "16-\nA- English  B- Englich\nC- Endlish  D- Emglish"
t.cell(0,2).text = "17-\nA- waiter  B- water\nC- woter  D- witer"
t.cell(0,3).text = "18-\nA- shef  B- Chef\nC- cheef  D- sheef"

add_para("")
add_bold_para("Writing:    /2")
add_line("Write these information in complete sentences: your [name / age / job / future job]:")
add_line("……………………………………………………………………………………………………")
add_line("……………………………………………………………………………………………………")

add_bold_para("Good luck! T. Zahrah", WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# EXAM 6: AboKhaled Amer - The Last Touch (Page 20)
# ============================================================
section_break()
add_heading_centered("Second Term 1447 H – First Session Test", level=1)
add_heading_centered("First Intermediate – The Last Touch", level=2)
add_bold_para("Creativity – AboKhaled Amer", WD_ALIGN_PARAGRAPH.CENTER)

add_line("الصف: الأول المتوسط          اسم الطالب: ……………………………………")

add_bold_para("Write the correct word under each picture: (6 M)")
add_line("Math – Art – two o'clock – teacher – waiter – eat – cat")
add_line("[7 pictures with blanks]")
add_line("[    ]    [    ]    [    ]    [    ]    [    ]    [    ]    [    ]")

add_separator()
add_bold_para("Match: (3M)")
t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
t.cell(0,0).text = "1. How old are you?"
t.cell(0,1).text = "I'm a student."
t.cell(1,0).text = "2. What's your favorite subject?"
t.cell(1,1).text = "I'm from Saudi Arabia."
t.cell(2,0).text = "3. Where are you from?"
t.cell(2,1).text = "My favorite subject is Math."
t.cell(3,0).text = "4. What do you want to be?"
t.cell(3,1).text = "I am 13 years old."
t.cell(4,0).text = "5. What do you do?"
t.cell(4,1).text = "I want to be a pilot."
t.cell(5,0).text = "6. What time is it?"
t.cell(5,1).text = "It's nine fifteen."
t.cell(6,0).text = ""
t.cell(6,1).text = "My name is Omar."

add_para("")
add_bold_para("Choose the correct answer: (6M)")
add_line("1. She ( go – goes – going ) to school by bus.")
add_line("2. What ( do – does – doing ) you do?")
add_line("3. They ( don't – doesn't – does ) speak French.")
add_line("4. He ( have – has – are ) black hair.")
add_line("5. I go to school ( in – on – at ) 7 o'clock.")
add_line("6. He is ( studying – study – studied ) in Riyadh this year.")

add_separator()
add_bold_para("Write the missing letter: (5M)")
add_line("( o – g – n – c – t )")
add_line("Geo__raphy    wa__ch    mecha__ic    Scien__e    doct__r")

add_bold_para("End of Questions", WD_ALIGN_PARAGRAPH.CENTER)

# Save
outpath = os.path.expanduser("~/Desktop/اختبار-انجليزي-اول-متوسط-ف2-1447.docx")
doc.save(outpath)
print(f"Saved: {outpath}")
