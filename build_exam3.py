#!/usr/bin/env python3
"""Build Word document for 3rd Intermediate English Exam - Second Term 1447"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = 'Times New Roman'
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

def hc(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def bp(text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p

def ap(text): return doc.add_paragraph(text)
def sep(): doc.add_paragraph("─" * 60)
def pb(): doc.add_page_break()

def grid(rows_data, cols):
    t = doc.add_table(rows=len(rows_data), cols=cols)
    t.style = 'Table Grid'
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            if j < cols:
                t.cell(i, j).text = val
    return t

# ============================================================
# EXAM 1: 1st Period Exam /20 (Pages 2-5)
# ============================================================
hc("1st Period Exam – 1447 – Second Term", 1)
hc("3rd Intermediate Grade – English Language     /20", 2)
ap("Student's Name: ……………………………………    Class: ……………")

bp("1- Read the following passage and answer the questions below:    /5 marks")
ap("Ali and his father went to a garage sale early on Saturday morning. The garage sale was near their house. They saw many old things like toys, books, and clothes. Ali looked at everything carefully. He found a small red toy car and showed it to his father. His father asked the seller for the price and bought the toy because it was cheap. Ali also saw old storybooks and helped his father choose one. Ali felt happy and enjoyed spending time with his father.")
ap("")
bp("A. Choose the correct answer:")
grid([
    ["1- Where did Ali and his father go?", "a- museum", "b- bank", "c- garage sale"],
    ["2- What day did they go to the garage sale?", "a- on Saturday", "b- on Monday", "c- on Sunday"],
    ["3- They saw many old things like ……, books, and clothes.", "a- toys", "b- football", "c- cars"],
], 4)
ap("")
bp("B. Put ( ✓ ) or ( X ):")
ap("1- The garage sale was near their house. (        )")
ap("2- The toy car was very expensive. (        )")

sep()
bp("2- Choose the correct answer:    /5 marks")
grid([
    ["1- She studied in London, ……. she?", "a- didn't", "b- aren't", "c- doesn't"],
    ["2- Sara should………… the doctor.", "a- saw", "b- seen", "c- see"],
    ["3- You must ………. at the traffic light.", "a- stops", "b- stopping", "c- stop"],
    ["4- Omar drives ……………….", "a- slow", "b- slowly", "c- slows"],
    ["5- Don't …………… your car there.", "a- parks", "b- park", "c- parking"],
], 4)

sep()
bp("3- Do as shown between brackets:    /2 marks")
bp("(Use adverb of manner)")
ap("a- They are fast runner.")
ap("………………………………………………………………………………………….")
bp("(Reorder)")
ap("a- Why / we / order / a pizza / don't / ?")
ap("………………………………………………………………………………………….")

pb()
bp("4- Write the correct word under each picture:    /4 marks")
ap("windshield wiper – hose – saw – trunk")
ap("[Picture 1: trunk]    [Picture 2: hose]    [Picture 3: saw]    [Picture 4: windshield wiper]")
ap("……………………    ……………………    ……………………    ……………………")

sep()
bp("5- Choose the correct answer:    /2 marks")
ap("1- Three children born at the same time: ( baby – child – triplets )")
ap("2- Study and research: ( investigate – trend – heredity )")

sep()
bp("6- Choose the missing letter:    /2 marks")
ap("1- She was wearing a bei…e dress. ( s – g – z )")
ap("2- You must obey the road …..igns. ( q – s – u )")

sep()
bp("7- Write the missing letter:")
ap("[Picture: teapot]    …..ot")
ap("[Picture: sword]    s…….ord")
bp("GOOD LUCK", WD_ALIGN_PARAGRAPH.CENTER)

# Answer Key
pb()
hc("نموذج الإجابة – 1st Period Exam /20", 1)
ap("1- A: 1-c (garage sale), 2-a (on Saturday), 3-a (toys)")
ap("   B: 1- ✓, 2- X")
ap("2- 1-a (didn't), 2-c (see), 3-c (stop), 4-b (slowly), 5-b (park)")
ap("3- a: They run fast. | b: Why don't we order a pizza?")
ap("4- trunk / hose / saw / windshield wiper")
ap("5- 1: triplets, 2: investigate")
ap("6- 1: g (beige), 2: s (signs)")
ap("7- pot (p), sword (w)")

# ============================================================
# EXAM 2: First Period Exam - Oral/Written 20 (Pages 6-9)
# ============================================================
pb()
hc("First Period Exam – Second Term", 1)
hc("Third Intermediate Grade", 2)

t = doc.add_table(rows=2, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.cell(0,0).text = "المجموع\nTotal"
t.cell(0,1).text = "تحريري"
t.cell(0,2).text = "شفهي"
t.cell(1,0).text = "20"
t.cell(1,1).text = "10"
t.cell(1,2).text = "10"
for row in t.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

ap("Name ……………………………………          Class ………………………………")

sep()
bp("Comprehension    /2")
bp("Read the following then answer the questions:")
ap("Bobby Shafran started a new semester at Sullivan Community College in New York. The previous semester Eddy Galland was a student there. The two teens looked exactly alike. Another student confused Bobby with Eddy, and then he realized that the boys were probably brothers. That student introduced Bobby to Eddy. The two boys found out that they were in fact twins, born at the same time. The boys were orphans and grew up in two different families.")

bp("A- Put (✓) for true or (X) for false:")
ap("1. The previous semester Eddy was a student at Sullivan Community College. (          )")
ap("2. The boys were orphans and grew up in the same family. (          )")

bp("B- Choose the right answer:")
grid([
    ["1- Sullivan Community College is in _____.", "London", "New York", "Washington", "Tokyo"],
    ["2- Bobby and Eddy were in fact ______.", "teachers", "friends", "twins", "cousins"],
], 5)

sep()
bp("Grammar    /2.5")
bp("Choose the right answer:")
grid([
    ["1- You live in Dammam, ___ ___?", "A. aren't you", "B. do you", "C. don't you", "D. doesn't he"],
    ["2- ______ you coming with us?", "A. Aren't", "B. Don't", "C. Haven't", "D. Isn't"],
    ["3- She didn't work in a school, ____ ____?", "A. didn't she", "B. did she", "C. doesn't she", "D. does she"],
    ["4- How did she walk? She walked _________.", "A. slow", "B. happy", "C. quick", "D. quickly"],
    ["5- The adverb form of good is _______.", "A. good", "B. will", "C. well", "D. going"],
], 5)

pb()
bp("Vocabulary    /2")
bp("A- Choose the best answer to fill in the blank:")
grid([
    ["1. Someone you don't know left a message on your _________.", "book", "cell phone", "chair", "school"],
    ["2. ______ is a child who has lost his parents.", "Orphan", "Twin", "Confuse", "Aunt"],
    ["3. What do you use to clean the living room?", "a sofa", "a ladder", "a vacuum cleaner", "a wall"],
    ["4- The _____ turns the car on.", "ignition", "headlight", "signal light", "GPS"],
], 5)

ap("")
bp("B- Put the right word under its picture:    /2.5")
ap("pliers – hammer – seat belt – hose – brake pedal")
ap("[Picture 1: brake pedal]  [Picture 2: seat belt]  [Picture 3: hose]  [Picture 4: pliers]  [Picture 5: hammer]")
ap("______________  ______________  ______________  ______________  ______________")

sep()
bp("Orthography    /1")
bp("Choose the correct letter:")
grid([
    ["1.", "Was Ahmed able to fi_ the car?", "l", "x", "k"],
    ["2.", "She found a set of nice crystal gla_ses", "s", "d", "v"],
], 5)

bp("Good luck          T.", WD_ALIGN_PARAGRAPH.CENTER)

# Answer Key
pb()
hc("نموذج الإجابة – First Period Exam (Oral/Written)", 1)
ap("Comprehension: A- 1: ✓, 2: X | B- 1: New York, 2: twins")
ap("Grammar: 1-C (don't you), 2-A (Aren't), 3-B (did she), 4-D (quickly), 5-C (well)")
ap("Vocabulary: A- 1: cell phone, 2: Orphan, 3: a vacuum cleaner, 4: ignition")
ap("           B- brake pedal / seat belt / hose / pliers / hammer")
ap("Orthography: 1: x (fix), 2: s (glasses)")

# ============================================================
# EXAM 3: Quiz (1) - Ms. Tahani Al-Ghamdi (Pages 10-13)
# ============================================================
pb()
hc("Quiz (1) – 1447 – Second Term", 1)
hc("3rd Intermediate Grade     /10", 2)
ap("Name: ……………………………    Class: ……………")

bp("I. Comprehension:")
bp("Read the passage to answer the questions below:")
ap("Ali has a scooter. His scooter is green. His uncle gave him the scooter on his 9th birthday as a present. Ali loves it so much! Every Saturday, Ali and his brother go to the park nearby their house. He rides the scooter while his brother rides his bike. They have so much fun at the park.")

bp("A- Choose the correct answer:")
ap("1- At the weekend, Ali goes to the ………………")
ap("    a- toy shop       b- park       c- school")
ap("2- On Ali's birthday, his …………… gave him a gift.")
ap("    a- mother       b- brother       c- uncle")

bp("B- Check ( T ) or ( F ):")
ap("1- Ali's brother has a scooter. (    )")
ap("2- Ali is 8 years old. (    )")

sep()
bp("II. Grammar:")
bp("A- Choose:")
ap("1- Let's (go – went – going) to the beach this afternoon.")
ap("2- She ate a burger, (don't she – doesn't she – didn't she)?")
ap("3- Mona is a good cook. She cooks (good – well – better).")
ap("4- You (must – mustn't – shouldn't) listen to your teacher.")

bp("B- Do as shown between brackets:")
ap("1- Why don't we have a garage sale?")
ap("…………………………………………………………. (Refuse the suggestion)")
ap("2- (Report the requests and commands)")
ap("   a- Nada: \"Clean the room\".")
ap("   ………………………………………………………….")

sep()
bp("III. Vocabulary:")
bp("A- Classify the words:")
ap("knife – vacuum cleaner – stop")
grid([
    ["Things used for cleaning", "Things used for cooking", "Traffic signs"],
    ["", "", ""],
], 3)

pb()
bp("B- Write the words under the pictures:")
ap("ignition – broom – rearview mirror – hammer")
ap("[Picture 1: rearview mirror]    [Picture 2: ignition]    [Picture 3: hammer]")
ap("______________       ______________       ______________")

sep()
bp("IV. Orthography:")
bp("A- Choose the correct missing letters:")
grid([
    ["1- They have a lot of anti__ue furniture.", "a- c", "b- k", "c- q"],
    ["2- We use a fo__k to eat food.", "a- l", "b- r", "c- c"],
    ["3- He was arrested for driving without a licen__e.", "a- c", "b- s", "c- z"],
    ["4- We must o__ey the road signs.", "a- b", "b- d", "c- p"],
], 4)

bp("Good Luck ☺", WD_ALIGN_PARAGRAPH.CENTER)
bp("Ms. Tahani Al-Ghamdi", WD_ALIGN_PARAGRAPH.CENTER)

# Answer Key
pb()
hc("نموذج الإجابة – Quiz (1) Ms. Tahani", 1)
ap("I. A- 1: b (park), 2: c (uncle) | B- 1: F, 2: F")
ap("II. A- 1: go, 2: didn't she, 3: well, 4: must")
ap("    B- 1: Sorry, I can't. | 2: Nada told me to clean the room.")
ap("III. A- Cleaning: vacuum cleaner | Cooking: knife | Traffic: stop")
ap("     B- rearview mirror / ignition / hammer")
ap("IV. 1: q (antique), 2: r (fork), 3: s (license), 4: b (obey)")

# ============================================================
# EXAM 4: Quiz (3) - Ms. Tahani / Hefar School (Pages 14-17)
# ============================================================
pb()
hc("Quiz (3) – Hefar Intermediate School", 1)
hc("3rd Intermediate Grade     /20", 2)
ap("Name: ……………………………    Class: ……………")

bp("I. Comprehension:")
bp("Use the following word bank to write a paragraph about what you are learning to do.")
ap("WORD BANK: (year – speak – alphabet – conversation – understand – fluently)")
ap("_______________________________________________________________")
ap("_______________________________________________________________")

sep()
bp("II. Grammar:")
bp("A- Choose:")
ap("1- Let's (go – went – going) to the beach this afternoon.")
ap("2- He will (was able to – were able to – be able to) late.")
ap("3- She ate a burger, (don't she – doesn't she – didn't she)?")

bp("B- Do as shown between brackets:")
ap("1- Why don't we have a garage sale?")
ap("…………………………………………………… (Accept the suggestion)")
ap("2- They aren't football players, ……………..? (Complete the tag question)")

sep()
bp("III. Vocabulary:")
bp("A- Write the words under the pictures:")
ap("saw – knife – teddy bear – luggage – hammer")
ap("[Picture 1: hammer]  [Picture 2: knife]  [Picture 3: saw]  [Picture 4: teddy bear]")
ap("______________  ______________  ______________  ______________")

bp("B- Fill in the blank with the correct words in the box:")
ap("heredity – lawn mower – talkative")
ap("1- My friend is a ……………… person.")
ap("2- We need a ……………………… to cut the grass.")

sep()
bp("IV. Orthography:")
bp("A- Choose the correct missing letters:")
grid([
    ["1- They have a lot of ant__que furniture.", "a- i", "b- o", "c- u"],
    ["2- We use a f__rk to eat food.", "a- u", "b- o", "c- a"],
    ["3- You shouldn't use the ho__e to wash your car.", "a- c", "b- s", "c- z"],
    ["4- The best way to open it is with p__iers.", "a- r", "b- n", "c- l"],
    ["5- He is drinking with a c__p and saucer.", "a- u", "b- a", "c- k"],
    ["6- She is sweeping with a b__oom.", "a- l", "b- r", "c- t"],
], 4)

bp("Good Luck ☺", WD_ALIGN_PARAGRAPH.CENTER)
bp("Ms. Tahani Al-Ghamdi", WD_ALIGN_PARAGRAPH.CENTER)

# Answer Key
pb()
hc("نموذج الإجابة – Quiz (3)", 1)
ap("I. Comprehension: Varied (free writing)")
ap("II. A- 1: go, 2: be able to, 3: didn't she")
ap("    B- 1: That's a good idea / OK / Sure. | 2: are they")
ap("III. A- hammer / knife / saw / teddy bear")
ap("     B- 1: talkative, 2: lawn mower")
ap("IV. 1: i (antique), 2: o (fork), 3: s (hose), 4: l (pliers), 5: u (cup), 6: r (broom)")

# ============================================================
# EXAM 5: Reading comp + Grammar + Vocab (Manal Al-Shareef) Pages 18-19
# ============================================================
pb()
hc("1st Period Exam – English Language", 1)
hc("3rd Intermediate Grade", 2)
ap("Name: ……………………………………    Class: ……………    Mark: ………………")

bp("❊ Reading Comprehension ❊    /5 marks")
ap("Mira and Lila were twin sisters who grew up in different families. They didn't meet until they were 42 years old.\nWhen they finally met, they were surprised—they were both wearing the same clothes! As they talked, they found more similarities. They both loved painting, drank hot tea with lemon, and their cats had the same name. They laughed a lot together, so people called them \"The Laughing Sisters.\"")

bp("1) Fill in the blank with the correct words:")
grid([
    ["1. Mira and Lila were:", "a. friends", "b. twin sisters", "c. cousins"],
    ["2. They grew up in:", "a. different family", "b. the same family", "c. the same city"],
    ["3. How old were they when they met?", "a. 20", "b. 32", "c. 42"],
], 4)

bp("2) Choose (✓) or (✗):")
ap("1. They both liked hot tea with lemon. (  )")
ap("2. They had dogs with the same name. (  )")

sep()
bp("❊ Grammar ❊    /7 marks")
bp("Choose the correct:")
grid([
    ["1. They aren't from here, ______they?", "are", "is", "aren't"],
    ["2. Ahmed should ______ healthy food.", "eat", "ate", "eating"],
    ["3. He brushed his teeth this morning, _____he?", "did", "don't", "didn't"],
    ["4. You ______ stop at the traffic lights.", "must", "should", "would"],
    ["5. My father is a hard worker. He works __________.", "hards", "hard", "hardly"],
], 4)

bp("Write the correct answer:")
ap("1. Tariq drives under the speed limit. He drives _________ (slow)")
ap("2. order / some / Let's / food ______________________________. (reorder)")

pb()
bp("❊ Vocabulary ❊    /6 marks")
bp("1) Match the following:")
grid([
    ["1. cleaning", "lights"],
    ["2. steering", "frying pan"],
    ["3. signal", "wheel"],
    ["4. cooking", "mop"],
], 2)

bp("2) Choose the correct answer:")
ap("1. A child who has lost his parents. [orphan / boy / kid]")
ap("2. Three children were born at the same time. [brothers / friends / triplets]")
ap("3. We use it to slow or stop the car. [dashboard / brake pedal / gas pedal]")
ap("4. ______ and broom are used for cleaning. [pot / fan / hose]")

bp("3) Write the correct words under each picture:")
ap("Windshield / Lawn mower / GPS / Luggage")
ap("[Picture 1: GPS]  [Picture 2: Lawn mower]  [Picture 3: Luggage]  [Picture 4: Windshield]")
ap("______________  ______________  ______________  ______________")

sep()
bp("❊ Orthography ❊    /2 marks")
bp("Choose the correct letter:")
ap("1. You should dri__e carefully. ( v – f – k )")
ap("2. We use a saw and ham__er for repairing. ( n – m – l )")
ap("3. You must obey the road si__ns. ( g – c – q )")
ap("4. I met my best friend by coi__cidence. ( m – h – n )")

bp("☺ With my best wishes ☺", WD_ALIGN_PARAGRAPH.CENTER)
bp("Focus with Manal          Teacher/Manal Al-Shareef", WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# EXAM 6: AboKhaled Amer - The Last Touch (Page 20)
# ============================================================
pb()
hc("Second Term 1447 H – First Session Test", 1)
hc("Third Intermediate – The Last Touch", 2)
bp("Creativity – AboKhaled Amer", WD_ALIGN_PARAGRAPH.CENTER)

ap("اسم الطالب: ……………………………………          الصف: الثالث المتوسط")

bp("Write the correct word under each picture: (6M)")
ap("hammer – No Parking – ladder – knife – Stop – gas tank – fish")
ap("[7 pictures with blanks]")
ap("[    ]    [    ]    [    ]    [    ]    [    ]    [    ]    [    ]")

sep()
bp("Match the following: (6M)")
grid([
    ["1. I'm doing 60.", "a. board"],
    ["2. steering", "b. a way to respond to an introduction."],
    ["3. dash", "c. wheel"],
    ["4. How do you do?", "d. a machine to cut lawn."],
    ["5. lawn mower", "e. I'm driving 60 kilometers per hour."],
    ["6. pot", "f. something which you use to cook food."],
    ["", "g. play football."],
], 2)

ap("")
bp("Choose the correct answer: (6M)")
ap("1. She doesn't come, ( doesn't she – she doesn't – does she – she does )?")
ap("2. He ( have – are – were – is ) able to clean the room.")
ap("3. Tomorrow, I ( was – am – will – have ) be able to come.")
ap("4. You ( must – mustn't – is must – should ) stop at the traffic lights.")
ap("5. You ( should – shouldn't – is should – must ) eat when you're driving.")
ap("6. She plays ( quiet – quietly – quietness – quieter ).")

sep()
bp("Supply the missing letter: (2M)    ( g – o – s – t )")
ap("[Picture: spoon]  spo__n")
ap("[Picture: seat belt]  sea__ belt")
ap("[Picture: hose]  ho__e")
ap("[Picture: danger sign]  dan__er")

bp("End of Questions", WD_ALIGN_PARAGRAPH.CENTER)

# Save
outpath = os.path.expanduser("~/Desktop/اختبار-انجليزي-ثالث-متوسط-ف2-1447.docx")
doc.save(outpath)
print(f"Saved: {outpath}")
