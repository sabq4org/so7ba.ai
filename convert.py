#!/usr/bin/env python3
import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re, os

os.chdir('/home/openclaw/.openclaw/workspace')

# Read markdown
with open('alwatan-report.md', 'r', encoding='utf-8') as f:
    md_text = f.read()

# --- DOCX ---
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
font.rtl = True

for line in md_text.split('\n'):
    line = line.rstrip()
    if not line:
        continue
    
    # Clean markdown bold/italic
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
    clean = re.sub(r'\*(.+?)\*', r'\1', clean)
    
    if line.startswith('# '):
        p = doc.add_heading(clean[2:], level=1)
    elif line.startswith('## '):
        p = doc.add_heading(clean[3:], level=2)
    elif line.startswith('### '):
        p = doc.add_heading(clean[4:], level=3)
    elif line.startswith('---'):
        doc.add_paragraph('─' * 50)
    elif line.startswith('- '):
        doc.add_paragraph(clean[2:], style='List Bullet')
    else:
        p = doc.add_paragraph()
        # Handle bold
        parts = re.split(r'(\*\*.+?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

doc.save('alwatan-report.docx')
print("DOCX created")

# --- HTML for PDF-like viewing ---
html = markdown.markdown(md_text, extensions=['extra'])
html_doc = f'''<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.8; direction: rtl; }}
h1 {{ color: #1a1a2e; border-bottom: 3px solid #16213e; padding-bottom: 10px; }}
h2 {{ color: #16213e; border-bottom: 1px solid #ddd; padding-bottom: 8px; margin-top: 40px; }}
h3 {{ color: #0f3460; }}
hr {{ border: none; border-top: 2px solid #e0e0e0; margin: 30px 0; }}
</style>
</head>
<body>{html}</body>
</html>'''

with open('alwatan-report.html', 'w', encoding='utf-8') as f:
    f.write(html_doc)
print("HTML created")
