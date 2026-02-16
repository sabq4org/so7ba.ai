from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# RTL support helper
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)

def set_cell_rtl(cell):
    for p in cell.paragraphs:
        set_rtl(p)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ===== TITLE =====
p = doc.add_paragraph()
set_rtl(p)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('تحليل إطار الحصانة الرقمية')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p2 = doc.add_paragraph()
set_rtl(p2)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('تقييم شامل — نقاط القوة والتطوير')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

# ===== SECTION: نقاط القوة =====
h = doc.add_heading('💪 نقاط القوة', level=1)
set_rtl(h)

strengths = [
    ('هيكل استراتيجي محكم', '8 محاور مترابطة تغطي الدورة الكاملة: من الوقاية → الاستجابة → الهجوم → القياس. التسلسل منطقي: ابنِ مناعة أولاً، ثم تعلّم الرد، ثم اهجم بذكاء.'),
    ('بروتوكول الاستجابة (0-72 ساعة)', 'من أقوى الأجزاء — تقسيم زمني واضح ومهني جداً. "قاعدة الأسئلة الثلاثة" = ذهب خالص — توفر طاقة وتمنع ردود متسرعة.'),
    ('الخطوط الحمراء واضحة', '❌ لا بوتات | ❌ لا شراء ترندات | ❌ لا تزييف — هذا يعطي مصداقية عالية للإطار نفسه.'),
    ('المبدأ الأساسي قوي جداً', '"المصداقية أقوى سلاح" | "الجودة تغلب الكمية" | "الدقة تغلب الكثرة"'),
    ('الميزانية واقعية ومفصلة', '2.25 مليون شهرياً مع توزيع واضح — 67 شخص إجمالي (32 فريق عمليات + 20 شبكة مصداقية + 15 محتوى).'),
    ('مؤشرات أداء حقيقية', 'تحول المشاعر 75% | جودة الوصول 85% | سرعة الاحتواء < 24 ساعة | مناعة 80%'),
]

for title, desc in strengths:
    p = doc.add_paragraph()
    set_rtl(p)
    run_t = p.add_run(f'✅ {title}')
    run_t.bold = True
    run_t.font.size = Pt(13)
    run_t.font.color.rgb = RGBColor(0x0D, 0x6E, 0x3A)
    p2 = doc.add_paragraph()
    set_rtl(p2)
    p2.add_run(desc).font.size = Pt(11)
    p2.paragraph_format.left_indent = Inches(0.3)

# ===== SECTION: نقاط تحتاج تطوير =====
doc.add_page_break()
h2 = doc.add_heading('⚠️ نقاط تحتاج تطوير', level=1)
set_rtl(h2)

weaknesses = [
    ('غياب السيناريوهات التطبيقية', 'الإطار نظري ممتاز، لكن ينقصه أمثلة عملية.', 'أضف 3-5 سيناريوهات واقعية (أزمة تغريدة، حملة منظمة، تسريب...).'),
    ('الأسئلة الثلاثة — ما ذُكرت!', 'القاعدة تقول "3 أسئلة" لكن ما وضّح وش هي بالضبط.', 'حددها صراحة (مثلاً: هل يؤثر؟ هل جمهورنا شافه؟ هل الرد يزيده انتشار؟).'),
    ('تكتيكات المنصات سطحية', 'ذكر X, TikTok, Snapchat, YouTube بدون تفصيل لكل منصة.', 'لكل منصة: نوع المحتوى + التوقيت + النبرة + أدوات القياس.'),
    ('غياب خطة التدريب', '32 شخص يحتاجون تدريب مستمر.', 'تمارين محاكاة ربع سنوية (War Room Simulation).'),
    ('الذكاء الاصطناعي غائب', 'ما في ذكر لأدوات AI في الرصد والتحليل.', 'أضف طبقة AI لـ: رصد المشاعر التلقائي + تحليل الحسابات المشبوهة + توليد تقارير فورية.'),
    ('خطة التصعيد غير واضحة', 'متى ننتقل من "استجابة" إلى "هجوم ذكي"؟', 'مصفوفة تصعيد واضحة بمعايير كمية.'),
]

for title, issue, suggestion in weaknesses:
    p = doc.add_paragraph()
    set_rtl(p)
    run_t = p.add_run(f'⚠️ {title}')
    run_t.bold = True
    run_t.font.size = Pt(13)
    run_t.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    
    p2 = doc.add_paragraph()
    set_rtl(p2)
    p2.add_run(f'المشكلة: {issue}').font.size = Pt(11)
    p2.paragraph_format.left_indent = Inches(0.3)
    
    p3 = doc.add_paragraph()
    set_rtl(p3)
    run_s = p3.add_run(f'💡 الاقتراح: {suggestion}')
    run_s.font.size = Pt(11)
    run_s.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    p3.paragraph_format.left_indent = Inches(0.3)

# ===== SECTION: جدول التقييم =====
doc.add_page_break()
h3 = doc.add_heading('📊 جدول التقييم الشامل', level=1)
set_rtl(h3)

table = doc.add_table(rows=9, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['المعيار', 'التقييم']
rows_data = [
    ('الرؤية الاستراتيجية', '⭐⭐⭐⭐⭐'),
    ('الهيكل التنظيمي', '⭐⭐⭐⭐⭐'),
    ('بروتوكول الاستجابة', '⭐⭐⭐⭐⭐'),
    ('التفصيل التنفيذي', '⭐⭐⭐'),
    ('تكتيكات المنصات', '⭐⭐'),
    ('أدوات التقنية والـ AI', '⭐⭐'),
    ('المصداقية والأخلاقيات', '⭐⭐⭐⭐⭐'),
    ('الميزانية والموارد', '⭐⭐⭐⭐'),
]

for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    set_cell_rtl(cell)

for idx, (criterion, rating) in enumerate(rows_data):
    table.rows[idx+1].cells[0].text = criterion
    table.rows[idx+1].cells[1].text = rating
    set_cell_rtl(table.rows[idx+1].cells[0])
    set_cell_rtl(table.rows[idx+1].cells[1])

doc.add_paragraph('')
p_score = doc.add_paragraph()
set_rtl(p_score)
p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_score = p_score.add_run('التقييم العام: 8 / 10')
run_score.bold = True
run_score.font.size = Pt(18)
run_score.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p_desc = doc.add_paragraph()
set_rtl(p_desc)
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_desc.add_run('إطار قوي جداً على المستوى الاستراتيجي، يحتاج تعزيز في التفاصيل التنفيذية والتقنية.').font.size = Pt(12)

# ===== LAST PAGE: خلاصة رأي صُحبة =====
doc.add_page_break()

# Background-colored box effect via table with shading
summary_table = doc.add_table(rows=1, cols=1)
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = summary_table.rows[0].cells[0]

# Set cell shading to dark blue
shading = cell._element.get_or_add_tcPr()
shading_elm = shading.makeelement(qn('w:shd'), {
    qn('w:val'): 'clear',
    qn('w:color'): 'auto',
    qn('w:fill'): '1B3A5C'
})
shading.append(shading_elm)

# Title
p_title = cell.paragraphs[0]
set_rtl(p_title)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run('🧠 خلاصة رأي صُحبة')
run_title.bold = True
run_title.font.size = Pt(22)
run_title.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p_line = cell.add_paragraph()
set_rtl(p_line)
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run('━' * 40)
run_line.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)

opinion_text = """هذا الإطار شغل محترف وواضح إنه مبني على خبرة حقيقية في إدارة الأزمات الرقمية.

النقطة الأقوى: فلسفة "المصداقية أقوى سلاح" + رفض البوتات والتزييف. هذا يخلّي الإطار يصلح كمرجع طويل المدى، مو مجرد خطة مؤقتة.

بروتوكول الـ 72 ساعة ممتاز ويصلح للتطبيق الفوري.

لكن بصراحة:
• الجانب التنفيذي يحتاج تفصيل أكثر — السيناريوهات العملية ناقصة
• تكتيكات المنصات سطحية — كل منصة لها لعبة مختلفة
• الذكاء الاصطناعي غائب تماماً وهذا فجوة كبيرة في 2026
• خطة التدريب غير موجودة — 32 شخص بدون تدريب = مخاطرة

هل يستحق التطبيق؟ نعم — بشرط إضافة:
1. ملاحق تنفيذية (سيناريوهات + قوالب ردود جاهزة)
2. طبقة ذكاء اصطناعي للرصد والتحليل
3. برنامج تدريب وتمارين محاكاة ربع سنوية

النتيجة المتوقعة عند التطبيق الكامل:
✅ تقليل زمن الاستجابة للأزمات بنسبة 60-70%
✅ تحسين السمعة الرقمية خلال 6-12 شهر
✅ بناء مناعة حقيقية ضد الحملات المنظمة
✅ عائد استثمار إيجابي خلال السنة الأولى

التقييم النهائي: 8/10 — أساس ممتاز يحتاج تعزيز تنفيذي وتقني."""

p_opinion = cell.add_paragraph()
set_rtl(p_opinion)
run_op = p_opinion.add_run(opinion_text)
run_op.font.size = Pt(12)
run_op.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Signature
p_sig = cell.add_paragraph()
set_rtl(p_sig)
p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_sig = p_sig.add_run('\n— صُحبة 🧠🤝')
run_sig.bold = True
run_sig.font.size = Pt(14)
run_sig.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)

out = '/home/openclaw/.openclaw/workspace/تحليل_إطار_الحصانة_الرقمية.docx'
doc.save(out)
print(f'DONE: {out}')
