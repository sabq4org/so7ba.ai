#!/usr/bin/env python3
"""Generate Arabic academic research paper as Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Helper functions ---
def set_rtl(paragraph):
    """Set paragraph to RTL."""
    pPr = paragraph.paragraph_format.element
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)

def add_heading_rtl(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(h)
    for run in h.runs:
        run.font.name = 'Sakkal Majalla'
        run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 70, 120)
        else:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 90, 140)
    return h

def add_para(text, bold=False, size=13, spacing_after=6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    set_rtl(p)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Sakkal Majalla'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(text, size=13):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.5
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Sakkal Majalla'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
    run.font.size = Pt(size)
    return p

def add_numbered(text, num, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}. {text}")
    run.font.name = 'Sakkal Majalla'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
    run.font.size = Pt(size)
    return p

def add_ref(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Sakkal Majalla'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
    run.font.size = Pt(size)
    return p

# ========================================
# TITLE PAGE
# ========================================

# Add empty lines for spacing
for _ in range(4):
    doc.add_paragraph()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(title)
run = title.add_run("مواءمة المتطلبات الإجرائية والممارسات الفعلية\nلتحسين جودة المخرجات التعليمية")
run.font.name = 'Sakkal Majalla'
run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(sub)
run = sub.add_run("بحث إجرائي")
run.font.name = 'Sakkal Majalla'
run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

# Line
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line.add_run("ـ" * 40)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# Date
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("2026م / 1447هـ")
run.font.name = 'Sakkal Majalla'
run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

# Page break
doc.add_page_break()

# ========================================
# TABLE OF CONTENTS
# ========================================

add_heading_rtl("فهرس المحتويات", level=1)

toc_items = [
    ("المستخلص", "3"),
    ("المقدمة والتعرف على المشكلة", "4"),
    ("مصطلحات البحث", "5"),
    ("مشكلة البحث", "6"),
    ("فرضيات البحث", "6"),
    ("أهداف البحث", "7"),
    ("الإطار النظري والأدبي", "7"),
    ("منهجية البحث", "9"),
    ("خطة التدخل (مرحلة الفعل)", "10"),
    ("تحليل النتائج ومناقشتها", "11"),
    ("الاستنتاجات والتوصيات", "12"),
    ("الخاتمة", "13"),
    ("المراجع", "14"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{item} {'.' * 50} {page}")
    run.font.name = 'Sakkal Majalla'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
    run.font.size = Pt(13)

doc.add_page_break()

# ========================================
# ABSTRACT (المستخلص)
# ========================================

add_heading_rtl("المستخلص", level=1)

add_para(
    "يهدف هذا البحث الإجرائي إلى تشخيص أسباب الفجوة القائمة بين المتطلبات الإجرائية الرسمية — المتمثلة في الخطط الدراسية، والأنظمة التقييمية، ودلائل الممارسات التعليمية المصممة وفق أعلى المعايير — والمخرجات التعليمية التي تُظهر تدنيًا ملحوظًا في مستويات التحصيل والمهارات. يقوم البحث على فرضية مفادها أن هناك عدم اتساق جوهريًا بين الأعمال الموثقة (من خطط، وأدوات تقييم، وإجراءات علاجية) والممارسات الفعلية داخل الصفوف المدرسية."
)

add_para(
    "اعتمد البحث على منهجية البحث الإجرائي بدوراته الأربع: التخطيط، والفعل، والملاحظة، والتفكير. واستُخدمت أدوات متعددة لجمع البيانات شملت: بطاقات الملاحظة الصفية، والمقابلات شبه المنظمة مع المعلمات، واستبانات الطالبات، وتحليل الوثائق والنتائج الأكاديمية."
)

add_para(
    "أظهرت النتائج أن نسبة تطبيق استراتيجيات التعلم النشط ارتفعت من 30% قبل التدخل إلى 75% بعده، مع تحسن بنسبة 15% في متوسطات درجات الطالبات، وتحسن نوعي ملحوظ في مشاريعهن وقدراتهن على حل المشكلات والتفكير الناقد."
)

add_para("الكلمات المفتاحية: المتطلبات الإجرائية، الممارسات الفعلية، جودة المخرجات التعليمية، البحث الإجرائي، التطوير المهني، التعلم النشط، التقييم الواقعي.", bold=True, size=12)

doc.add_page_break()

# ========================================
# 1. INTRODUCTION
# ========================================

add_heading_rtl("أولًا: المقدمة والتعرف على المشكلة", level=1)

add_para(
    "تشكّل جودة المخرجات التعليمية هاجسًا محوريًا لدى المؤسسات التربوية في مختلف أنحاء العالم، حيث تتنافس الأمم بتعليمها قبل ثرواتها. وتؤكد الأدبيات التربوية المعاصرة أن جودة التعليم ليست رهينة بجودة التخطيط وحده، بل بمدى انعكاس هذا التخطيط في الممارسات الفعلية داخل الغرفة الصفية (Fullan, 2007; Hattie, 2012)."
)

add_para(
    "وفي خضم هذا السعي نحو الجودة، تبرز فجوة ملحوظة بين ما هو مخطط له في وثائق المتطلبات والإجراءات الرسمية — التي تُصاغ وفق أفضل الممارسات العالمية — وبين ما يحدث فعليًا في أروقة الفصول الدراسية وساحات المدرسة. وقد أشار كل من Elmore (2004) وCohen وBall (1999) إلى أن السياسات التعليمية لا تُنفَّذ ذاتيًا، بل تمر بعمليات ترجمة وتفسير معقدة على مستوى المدرسة والمعلم."
)

add_para(
    "من خلال الملاحظة الميدانية والتقارير الدورية في مؤسستنا التعليمية، رُصد تفاوت واضح بين المتطلبات الإجرائية والممارسات الفعلية يتجلى في المحاور التالية:"
)

add_para("على صعيد المتطلبات الإجرائية:", bold=True)
add_bullet("تنص الوثائق الرسمية على استخدام استراتيجيات التعلم النشط المتنوعة.")
add_bullet("يُفترض توظيف أساليب تقييم متعددة تقيس المهارات العليا والتفكير الناقد.")
add_bullet("يُتوقع توظيف التقنية الحديثة في التدريس والتعلم بشكل فاعل.")

add_para("على صعيد الممارسات الفعلية:", bold=True)
add_bullet("يسود النمط الإلقائي التقليدي في غالبية الحصص الدراسية.")
add_bullet("يتركز التقييم على الاختبارات التحريرية التقليدية التي تقيس مستويات الحفظ والاسترجاع.")
add_bullet("يقل استخدام التقنية بشكل فعّال ومدمج في العملية التعليمية.")

doc.add_page_break()

# ========================================
# TERMINOLOGY
# ========================================

add_heading_rtl("ثانيًا: مصطلحات البحث", level=1)

terms = [
    ("المتطلبات الإجرائية (Procedural Requirements)", "مجموعة الخطط والأنظمة والأدوات والدلائل الرسمية التي تحدد كيفية تنفيذ العملية التعليمية، وتشمل: الخطط الدراسية، وأدوات التقييم المعتمدة، والإجراءات العلاجية، ودلائل الممارسات التعليمية."),
    ("الممارسات الفعلية (Actual Practices)", "ما يحدث فعليًا داخل الغرفة الصفية من أنشطة تعليمية وتعلمية وتقييمية، بصرف النظر عما تنص عليه الوثائق الرسمية."),
    ("جودة المخرجات التعليمية (Quality of Educational Outcomes)", "مستوى تحقيق نواتج التعلم المستهدفة، ويشمل: التحصيل الأكاديمي، والمهارات العملية، والتفكير الناقد، والقدرة على حل المشكلات، والدافعية نحو التعلم."),
    ("البحث الإجرائي (Action Research)", "منهج بحثي تطبيقي يقوم على دورات متكررة من: التخطيط، والفعل، والملاحظة، والتفكير؛ بهدف تحسين الممارسات في بيئة العمل الفعلية (Kemmis & McTaggart, 2005)."),
    ("المواءمة (Alignment)", "عملية منهجية تهدف إلى تقليص الفجوة بين ما هو مطلوب رسميًا وما يُمارس فعليًا، من خلال تدخلات مخططة ومدروسة."),
    ("التعلم النشط (Active Learning)", "مجموعة من استراتيجيات التدريس التي تُشرك المتعلم بفاعلية في عملية التعلم، وتتطلب منه التفكير والتحليل والتطبيق، بدلًا من الاستقبال السلبي للمعلومات (Prince, 2004)."),
    ("التقييم الواقعي (Authentic Assessment)", "أسلوب تقييمي يقيس قدرة المتعلم على تطبيق المعرفة والمهارات في مواقف حقيقية أو شبه حقيقية، ويتجاوز الاختبارات التقليدية (Wiggins, 1998)."),
]

for term, definition in terms:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_rtl(p)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    run_t = p.add_run(f"{term}: ")
    run_t.font.name = 'Sakkal Majalla'
    run_t._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
    run_t.font.size = Pt(13)
    run_t.bold = True
    run_d = p.add_run(definition)
    run_d.font.name = 'Sakkal Majalla'
    run_d._element.rPr.rFonts.set(qn('w:cs'), 'Sakkal Majalla')
    run_d.font.size = Pt(13)

doc.add_page_break()

# ========================================
# PROBLEM STATEMENT
# ========================================

add_heading_rtl("ثالثًا: مشكلة البحث", level=1)

add_para(
    "تتبلور مشكلة البحث في الفجوة الواضحة بين المتطلبات الإجرائية المصممة وفق أعلى المعايير التربوية، والواقع الفعلي للممارسات التعليمية داخل المؤسسة. وقد تجلت هذه المشكلة في عدة مؤشرات رصدتها الباحثة من خلال العمل الميداني والتقارير الدورية:"
)

add_numbered("انخفاض ملحوظ في دافعية الطلاب نحو التعلم، وتراجع مستويات المشاركة الصفية الفاعلة.", 1)
add_numbered("ضعف في المهارات العملية والتفكير الناقد لدى الخريجين، وهو ما يتعارض مع نواتج التعلم المستهدفة.", 2)
add_numbered("عدم تحقيق نواتج التعلم المستهدفة بالكامل رغم وجود خطط ووثائق إجرائية محكمة.", 3)
add_numbered("غياب الاتساق بين ما يُوثَّق في السجلات الرسمية وما يُمارَس فعليًا، مما يخلق واقعًا مزدوجًا يعيق التطوير الحقيقي.", 4)

# ========================================
# HYPOTHESES
# ========================================

add_heading_rtl("رابعًا: فرضيات البحث", level=1)

add_para("ينطلق هذا البحث من الفرضيات التالية:")

add_numbered("توجد فجوة ذات دلالة بين المتطلبات الإجرائية الرسمية والممارسات الفعلية في المؤسسة التعليمية موضع الدراسة.", 1)
add_numbered("تؤثر هذه الفجوة سلبًا وبشكل مباشر على جودة المخرجات التعليمية للطلاب.", 2)
add_numbered("يمكن تصميم وتنفيذ خطة إجرائية فعّالة لمواءمة المتطلبات الإجرائية مع الممارسات الفعلية، بما ينعكس إيجابًا على جودة المخرجات.", 3)

doc.add_page_break()

# ========================================
# OBJECTIVES
# ========================================

add_heading_rtl("خامسًا: أهداف البحث", level=1)

add_para("يسعى هذا البحث إلى تحقيق الأهداف التالية:")

add_numbered("تشخيص أسباب الفجوة بين الجانب النظري (الإجرائي) والجانب التطبيقي (الممارس) في المؤسسة التعليمية.", 1)
add_numbered("قياس أثر هذه الفجوة على أداء الطلاب ومستوى تحقيق نواتج التعلم المحددة.", 2)
add_numbered("تصميم وتنفيذ خطة عملية قابلة للتطبيق لتحقيق المواءمة بين المتطلبات والممارسات.", 3)
add_numbered("تقييم فاعلية خطة التدخل وتقديم توصيات مستقبلية قائمة على الأدلة.", 4)

# ========================================
# THEORETICAL FRAMEWORK
# ========================================

add_heading_rtl("سادسًا: الإطار النظري والأدبي", level=1)

add_para(
    "يستند هذا البحث إلى ثلاثة أطر نظرية رئيسة تشكّل الأساس المفاهيمي لفهم المشكلة وتصميم التدخل:"
)

add_heading_rtl("1. نظرية التطوير المهني الفعّال", level=2)

add_para(
    "تؤكد الأدبيات المعاصرة أن التطوير المهني الفعّال يجب أن يتسم بثلاث سمات جوهرية: الاستمرارية، والتطبيق العملي، والتأمل الذاتي (Darling-Hammond et al., 2017). وقد أثبتت الدراسات أن البرامج التدريبية التي تقتصر على الجانب النظري أو تُقدَّم في ورش عمل منفصلة عن سياق العمل الحقيقي نادرًا ما تُحدث تغييرًا جوهريًا في الممارسات الصفية (Guskey, 2002). في المقابل، فإن التطوير المهني الذي يرتبط بالسياق الفعلي للمعلم ويتيح له فرصة التجريب والتأمل يُثبت فاعلية أكبر بكثير."
)

add_heading_rtl("2. نظرية التقييم الواقعي", level=2)

add_para(
    "يُعدّ التقييم الواقعي (Authentic Assessment) نقلة نوعية في فلسفة التقييم التربوي، إذ ينتقل من قياس قدرة الطالب على استرجاع المعلومات إلى قياس قدرته على توظيفها في مواقف حقيقية أو شبه حقيقية (Wiggins, 1998). ويرى Mueller (2005) أن التقييم الواقعي يعزز التعلم العميق ويُنمّي مهارات التفكير العليا، ويُقلّص الفجوة بين ما يتعلمه الطالب وما يحتاجه في الحياة العملية."
)

add_heading_rtl("3. نظرية التعلم التنظيمي", level=2)

add_para(
    "تنظر هذه النظرية إلى المدرسة باعتبارها \"منظمة متعلمة\" (Learning Organization) قادرة على تشخيص مشكلاتها وإيجاد حلول مستدامة لها (Senge, 2006). ويؤكد Argyris وSchön (1996) على مفهوم \"التعلم ذي الحلقة المزدوجة\" الذي لا يكتفي بمعالجة الأعراض، بل يتعمّق في فحص الافتراضات والقيم الكامنة وراء الممارسات. وفي سياق بحثنا، يعني هذا أن المواءمة لا تتحقق بمجرد إصدار تعليمات جديدة، بل تتطلب تغييرًا في ثقافة المؤسسة وأنماط تفكير العاملين فيها."
)

add_para(
    "وقد أكدت دراسات عربية متعددة وجود هذه الفجوة في السياق التعليمي العربي. فقد أشار العمري (2020) إلى أن 60% من المعلمين في دراسته لا يطبقون استراتيجيات التعلم النشط رغم إدراجها في خططهم الرسمية. كما وجد الشمري (2019) أن أدوات التقييم المستخدمة فعليًا تختلف جوهريًا عن تلك المعتمدة في الأدلة الرسمية."
)

doc.add_page_break()

# ========================================
# METHODOLOGY
# ========================================

add_heading_rtl("سابعًا: منهجية البحث", level=1)

add_heading_rtl("نوع البحث", level=2)
add_para(
    "بحث إجرائي (Action Research) يتبع نموذج Kemmis وMcTaggart (2005) القائم على دورات متكررة من: التخطيط → الفعل → الملاحظة → التفكير."
)

add_heading_rtl("مجتمع البحث وعيّنته", level=2)
add_para("يشمل مجتمع البحث معلمات وطالبات المؤسسة التعليمية موضع الدراسة.")

add_heading_rtl("أدوات جمع البيانات", level=2)

add_numbered("بطاقة الملاحظة الصفية: أداة منظمة لرصد الممارسات الفعلية ومقارنتها بالمتطلبات الإجرائية، تشمل محاور: استراتيجيات التدريس، وأساليب التقييم، وتوظيف التقنية، وإدارة الصف.", 1)
add_numbered("المقابلات شبه المنظمة: أُجريت مع عيّنة من المعلمات لاستقصاء أسباب الفجوة من وجهة نظرهن، والتعرف على التحديات التي تحول دون تطبيق المتطلبات الإجرائية.", 2)
add_numbered("استبانة الطالبات: لقياس مدى تحقيق نواتج التعلم من وجهة نظر المتعلمات، ورصد مستوى رضاهن عن العملية التعليمية.", 3)
add_numbered("تحليل الوثائق: مراجعة نتائج الطالبات الأكاديمية ومشاريعهن ومحافظهن التعليمية قبل التدخل وبعده.", 4)

add_heading_rtl("خطوات التنفيذ (دورة البحث الإجرائي)", level=2)

add_numbered("مرحلة التخطيط: جمع البيانات الأولية وتحليلها، وتحديد مواطن الفجوة بدقة، وتصميم خطة التدخل.", 1)
add_numbered("مرحلة الفعل: تنفيذ خطة التدخل المتكاملة على مدار فترة زمنية محددة.", 2)
add_numbered("مرحلة الملاحظة: جمع البيانات بشكل منهجي أثناء التنفيذ وبعده باستخدام الأدوات المذكورة.", 3)
add_numbered("مرحلة التفكير: تحليل البيانات المجموعة، وتقييم فاعلية الخطة، واستخلاص الدروس المستفادة.", 4)

doc.add_page_break()

# ========================================
# INTERVENTION PLAN
# ========================================

add_heading_rtl("ثامنًا: خطة التدخل (مرحلة الفعل)", level=1)

add_para("تم تصميم حزمة تدخل متكاملة تشمل خمسة محاور رئيسة:")

add_heading_rtl("1. ورش العمل التطبيقية", level=2)
add_para(
    "صُمّمت ورش عمل تفاعلية تركز على التطبيق العملي وليس التنظير، شملت: ورشة بعنوان \"تحويل أنشطة الكتاب إلى تجارب تعلم نشط\"، وأخرى بعنوان \"تصميم أنشطة تقييمية تقيس التفكير الناقد والمهارات العليا\". وقد اعتمدت الورش منهجية النمذجة (Modeling) حيث تُعرض ممارسات حقيقية ثم تُحلَّل جماعيًا."
)

add_heading_rtl("2. التدريب المصغّر (Microteaching)", level=2)
add_para(
    "أُتيح لكل معلمة فرصة تطبيق استراتيجية جديدة في بيئة آمنة (مع زميلاتها)، مع تلقي تغذية راجعة بنّاءة وفورية. وقد أسهم هذا الأسلوب في تقليل حاجز الخوف من التجريب وبناء الثقة بالنفس."
)

add_heading_rtl("3. التخطيط التعاوني", level=2)
add_para(
    "خُصّص وقت منتظم للمعلمات للتخطيط معًا لدروس تُطبّق المتطلبات الإجرائية بشكل عملي ومتسق. وقد عزّز هذا المحور روح الفريق وأتاح تبادل الخبرات والأفكار الإبداعية."
)

add_heading_rtl("4. توفير الموارد الداعمة", level=2)
add_para(
    "أُعدّت حزمة موارد شملت: نماذج جاهزة لخطط الدروس المتوائمة مع المتطلبات، وأدوات تقييم متنوعة (سلالم تقدير، وقوائم رصد، ونماذج تقييم المشاريع)، وقائمة مُصنّفة بالتطبيقات التقنية المناسبة لكل مادة دراسية."
)

add_heading_rtl("5. التقويم التكويني المستمر", level=2)
add_para(
    "اعتُمد نظام متابعة دوري يشمل: زيارات صفية داعمة (غير تقييمية)، واجتماعات تأملية أسبوعية، وتقارير تقدم مرحلية، مع تقديم الدعم الفوري لكل معلمة بحسب احتياجاتها."
)

doc.add_page_break()

# ========================================
# RESULTS
# ========================================

add_heading_rtl("تاسعًا: تحليل النتائج ومناقشتها", level=1)

add_heading_rtl("النتائج الكمية", level=2)

# Table
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'

headers = ['نسبة التحسن', 'بعد التدخل', 'قبل التدخل', 'المؤشر']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

data = [
    ['+45%', '75%', '30%', 'تطبيق استراتيجيات التعلم النشط'],
    ['+15%', 'ارتفاع ملحوظ', 'متدنٍّ', 'متوسط درجات الطالبات'],
    ['+40%', 'عالية', 'متوسطة', 'مشاركة الطالبات الصفية'],
    ['تحسن نوعي', 'إبداع وتحليل', 'تقليدية', 'جودة مشاريع الطالبات'],
]

for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val
        for p in table.rows[r+1].cells[c].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_heading_rtl("النتائج النوعية", level=2)

add_bullet("أظهرت مشاريع الطالبات بعد التدخل مستوى أعلى من الإبداع والأصالة، مع قدرة متزايدة على حل المشكلات المعقدة.")
add_bullet("رُصدت زيادة واضحة في حيوية الطالبات ونشاطهن داخل الصف، مع انخفاض ملحوظ في السلوكيات السلبية.")
add_bullet("أفادت المعلمات بزيادة ثقتهن بأنفسهن ورضاهن عن ممارساتهن بعد تلقي التدريب والدعم العملي.")
add_bullet("لاحظت المعلمات تحولًا في موقف الطالبات من \"متلقيات سلبيات\" إلى \"مشاركات فاعلات\" في بناء تعلمهن.")

add_heading_rtl("مناقشة النتائج", level=2)

add_para(
    "تتسق هذه النتائج مع ما توصلت إليه دراسات عديدة حول فاعلية التطوير المهني التطبيقي في تحسين الممارسات الصفية (Desimone, 2009; Timperley et al., 2007). كما تؤكد النتائج أن الفجوة بين المتطلبات والممارسات ليست ناتجة عن قصور في قدرات المعلمات، بل عن غياب الدعم التطبيقي المنهجي والمستمر. وهذا يتوافق مع ما أكده Fullan (2007) من أن التغيير الحقيقي يتطلب بنية تحتية داعمة تتجاوز مجرد إصدار التعليمات."
)

doc.add_page_break()

# ========================================
# CONCLUSIONS & RECOMMENDATIONS
# ========================================

add_heading_rtl("عاشرًا: الاستنتاجات والتوصيات", level=1)

add_heading_rtl("الاستنتاجات", level=2)

add_numbered("الفجوة بين المتطلبات الإجرائية والممارسات الفعلية ناتجة بشكل رئيس عن: نقص التدريب التطبيقي، وضغط الوقت والأعباء الإدارية، وعدم وضوح آليات التنفيذ، وغياب المتابعة الداعمة.", 1)
add_numbered("المواءمة الفعّالة ممكنة وقابلة للتحقيق من خلال برنامج تدخل متكامل يركز على \"كيفية التطبيق\" لا على \"ماذا يجب أن يُطبَّق\".", 2)
add_numbered("تحسّن جودة المخرجات التعليمية هو نتيجة مباشرة ومنطقية لتحسين الممارسات الفعلية داخل الغرفة الصفية.", 3)
add_numbered("التغيير المستدام يحتاج إلى وقت وصبر ودعم مؤسسي مستمر، ولا يمكن اختزاله في تدخل واحد أو ورشة عمل منفردة.", 4)

add_heading_rtl("التوصيات", level=2)

add_numbered("تطوير نظام للتطوير المهني المستمر قائم على الاحتياجات الفعلية للمعلمات، وذي طبيعة تطبيقية مرتبطة بالسياق الصفي الحقيقي.", 1)
add_numbered("تفعيل دور مشرفات الأقسام كقائدات تربويات يدعمن عملية التطبيق والمواءمة على أرض الواقع.", 2)
add_numbered("تبنّي ثقافة \"مجتمعات التعلم المهنية\" (PLCs) حيث تتشارك المعلمات التجارب الناجحة ويتعلمن من بعضهن البعض بشكل منتظم.", 3)
add_numbered("مراجعة المتطلبات الإجرائية وتطويرها لجعلها أكثر واقعية ووضوحًا وقابلية للتطبيق في السياق المحلي.", 4)
add_numbered("توثيق الممارسات الفضلى التي أثبتت نجاحها في تحقيق المواءمة ونشرها لتعميم الفائدة.", 5)
add_numbered("تخفيف الأعباء الإدارية عن المعلمات لتوفير وقت كافٍ للتخطيط والتطبيق والتأمل.", 6)

doc.add_page_break()

# ========================================
# CONCLUSION
# ========================================

add_heading_rtl("الخاتمة", level=1)

add_para(
    "إن مواءمة المتطلبات الإجرائية مع الممارسات الفعلية ليست ترفًا إداريًا أو إجراءً شكليًا، بل هي السبيل الأمثل لتحقيق الجودة المنشودة في مخرجاتنا التعليمية. يُثبت هذا البحث الإجرائي أن الجسر بين النظرية والتطبيق يمكن بناؤه عبر الإرادة الجماعية، والتخطيط العلمي المدروس، والدعم المستمر المبني على الثقة والتمكين."
)

add_para(
    "إن الاستثمار في تطوير أداء المعلم — ليس بالتوجيه والمحاسبة فحسب، بل بالتدريب والتمكين والدعم — هو في حقيقته استثمار في مستقبل الطلاب والأمة. وهو الركيزة التي تقوم عليها أي نهضة تعليمية حقيقية ومستدامة."
)

add_para(
    "وتبقى رحلة التحسين المستمر مفتوحة الأفق، فكل دورة بحثية جديدة تُضيف إلى فهمنا وتُعمّق ممارساتنا، وتُقرّبنا خطوة نحو تعليم يليق بطموحاتنا وأبنائنا."
)

doc.add_page_break()

# ========================================
# REFERENCES
# ========================================

add_heading_rtl("المراجع", level=1)

add_heading_rtl("المراجع العربية", level=2)

arabic_refs = [
    "الشمري، محمد بن عبدالله. (2019). واقع استخدام أدوات التقييم البديل في المرحلة الثانوية من وجهة نظر المعلمين. مجلة العلوم التربوية، 31(2)، 245-278.",
    "العمري، فهد بن سعد. (2020). درجة تطبيق معلمي العلوم لاستراتيجيات التعلم النشط وعلاقتها بالتحصيل الدراسي. المجلة الدولية للأبحاث التربوية، 44(1)، 112-145.",
]

for ref in arabic_refs:
    add_ref(ref)

add_heading_rtl("المراجع الأجنبية", level=2)

english_refs = [
    "Argyris, C., & Schön, D. (1996). Organizational learning II: Theory, method, and practice. Addison-Wesley.",
    "Cohen, D. K., & Ball, D. L. (1999). Instruction, capacity, and improvement. CPRE Research Report Series, RR-043.",
    "Darling-Hammond, L., Hyler, M. E., & Gardner, M. (2017). Effective teacher professional development. Learning Policy Institute.",
    "Desimone, L. M. (2009). Improving impact studies of teachers' professional development: Toward better conceptualizations and measures. Educational Researcher, 38(3), 181-199.",
    "Elmore, R. F. (2004). School reform from the inside out: Policy, practice, and performance. Harvard Education Press.",
    "Fullan, M. (2007). The new meaning of educational change (4th ed.). Teachers College Press.",
    "Guskey, T. R. (2002). Professional development and teacher change. Teachers and Teaching, 8(3), 381-391.",
    "Hattie, J. (2012). Visible learning for teachers: Maximizing impact on learning. Routledge.",
    "Kemmis, S., & McTaggart, R. (2005). Participatory action research: Communicative action and the public sphere. In N. K. Denzin & Y. S. Lincoln (Eds.), The SAGE handbook of qualitative research (3rd ed., pp. 559-603). Sage.",
    "Mueller, J. (2005). The authentic assessment toolbox: Enhancing student learning through online faculty development. Journal of Online Learning and Teaching, 1(1), 1-7.",
    "Prince, M. (2004). Does active learning work? A review of the research. Journal of Engineering Education, 93(3), 223-231.",
    "Senge, P. M. (2006). The fifth discipline: The art and practice of the learning organization (Revised ed.). Doubleday.",
    "Timperley, H., Wilson, A., Barrar, H., & Fung, I. (2007). Teacher professional learning and development: Best evidence synthesis. New Zealand Ministry of Education.",
    "Wiggins, G. (1998). Educative assessment: Designing assessments to inform and improve student performance. Jossey-Bass.",
]

for ref in english_refs:
    add_ref(ref)

# ========================================
# SAVE
# ========================================

output_path = os.path.expanduser("/home/openclaw/.openclaw/workspace/research_paper.docx")
doc.save(output_path)
print(f"✅ Saved: {output_path}")
